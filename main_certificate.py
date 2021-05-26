#!/usr/bin/env python3

import os
from certificate import *
from docx import Document
import csv
# from docx2pdf import convert


# create output folder if not exist
try:
    os.makedirs("Output/Doc")
    os.makedirs("Output/PDF")
except OSError:
    pass

def get_participants(f):
    data = [] # create empty list
    with open(f, mode="r") as file:
        csv_reader = csv.DictReader(file)
        for row in csv_reader:
            data.append(row) # append all results
    return data

def create_docx_files(filename, list_participate, event, ambassador):

    for participate in list_participate:
        name = participate["Name Surname"]
        # use original file everytime
        doc = Document(filename)

        replace_participant_name(doc, name)
        replace_event_name(doc, event)
        replace_ambassador_name(doc, ambassador)
        doc.save('Output/Doc/{}.docx'.format(name))
        # convert('Output/Doc/{}.docx'.format(name), 'Output/Pdf/{}.pdf'.format(name))
        
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        print(p.text)

    

certificate_file = "Data Template/Event Certificate Template.docx"
participate_file = "Data Template/Event Participate Template.csv"

list_participate = get_participants(participate_file);
create_docx_files(certificate_file, list_participate, "WSL and VSCode ha", "Aysem Yas")



