#!/usr/bin/env python3
import re
import os
from docx import Document

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            print(p.text)
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


def replace_participant_name(doc, name):
    string = "Samplefirstname Samplelastname"
    reg = re.compile(r""+string)
    replace = r""+name
    docx_replace_regex(doc, reg , replace)

def replace_event_name(doc, event):
    string = "{INSERT EVENT NAME}"
    reg = re.compile(r""+string)
    replace = r""+event
    docx_replace_regex(doc, reg , replace)

def replace_ambassador_name(doc, name):
    
    reg = re.compile(r"STUDENT AMBASSADOR NAME")
    replace = r""+name
    docx_replace_regex(doc, reg , replace)

try:
    os.mkdir("Output")
except OSError:
    pass

filename = "Event Certificate Template.docx"
doc = Document(filename)
#STUDENT AMBASSADOR NAME 
replace_participant_name(doc, "Muhammed Oğuz Oğuz Oğuz Oğuz Oğuz")
replace_event_name(doc, "WSL + VSCode Edu HD Download")
replace_ambassador_name(doc, "Ayşem Aydoğan")
doc.save('Output/result1.docx')